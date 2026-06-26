"""
Convert Smart Classroom Markdown Documentation to DOCX
======================================================
Reads the markdown file and produces a professionally formatted .docx
"""

import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os
import zlib
import base64
import json
import urllib.request

INPUT_FILE = r"C:\Users\M.Yaswanth\.gemini\antigravity\brain\30aa54b5-df37-4ea0-8501-96113e9431f1\smart_classroom_documentation.md"
OUTPUT_FILE = r"C:\Data Science and Gen Ai projects\smart_classroom\Smart_Classroom_Documentation.docx"


def get_mermaid_image_path(mermaid_code, index):
    """Download diagram as PNG from mermaid.ink and save locally."""
    # Ensure images directory exists
    os.makedirs('images', exist_ok=True)
    img_path = f'images/diagram_{index}.png'
    
    # If already downloaded, reuse it
    if os.path.exists(img_path):
        return img_path
        
    try:
        data = {
            "code": mermaid_code,
            "mermaid": {
                "theme": "default"
            }
        }
        json_data = json.dumps(data).encode('utf-8')
        compressed = zlib.compress(json_data)
        b64_encoded = base64.urlsafe_b64encode(compressed).decode('ascii').rstrip("=")
        url = f"https://mermaid.ink/img/pako:{b64_encoded}"
        
        req = urllib.request.Request(url, headers={'User-Agent': 'Mozilla/5.0'})
        with urllib.request.urlopen(req, timeout=30) as response:
            with open(img_path, 'wb') as f:
                f.write(response.read())
        print(f"Downloaded diagram {index} successfully.")
        return img_path
    except Exception as e:
        print(f"Error downloading diagram {index}: {str(e)}")
        return None


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_formatted_text(paragraph, text):
    """Parse inline markdown formatting and add runs to paragraph."""
    # Process inline formatting: **bold**, `code`, [text](url), emojis
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`|\[.*?\]\(.*?\))', text)
    for part in parts:
        if not part:
            continue
        # Bold
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        # Inline code
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        # Link [text](url)
        elif re.match(r'\[.*?\]\(.*?\)', part):
            m = re.match(r'\[(.*?)\]\((.*?)\)', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
                run.underline = True
        else:
            paragraph.add_run(part)


def parse_table(lines):
    """Parse markdown table lines into list of rows (each row is list of cells)."""
    rows = []
    for line in lines:
        line = line.strip()
        if line.startswith('|') and line.endswith('|'):
            # Skip separator lines like |---|---|
            if re.match(r'^\|[\s\-:]+\|$', line.replace('|', '|').replace('-', '-')):
                continue
            cells = [c.strip() for c in line.split('|')[1:-1]]
            rows.append(cells)
    return rows


def create_docx(md_content):
    doc = Document()
    
    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # ── Define styles ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15
    
    for level in range(1, 5):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.bold = True
        if level == 1:
            hs.font.size = Pt(24)
            hs.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
        elif level == 2:
            hs.font.size = Pt(18)
            hs.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)
        elif level == 3:
            hs.font.size = Pt(14)
            hs.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
        elif level == 4:
            hs.font.size = Pt(12)
            hs.font.color.rgb = RGBColor(0x4A, 0x14, 0x8C)
    
    lines = md_content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    code_lang = ''
    in_mermaid = False
    mermaid_lines = []
    mermaid_index = 1
    
    while i < len(lines):
        line = lines[i]
        
        # ── Fenced code blocks (``` or ```language) ──
        if line.strip().startswith('```') and not in_code_block and not in_mermaid:
            lang = line.strip()[3:].strip().lower()
            if lang == 'mermaid':
                in_mermaid = True
                mermaid_lines = []
                i += 1
                continue
            in_code_block = True
            code_lang = lang
            code_lines = []
            i += 1
            continue
        
        if in_mermaid:
            if line.strip() == '```':
                # Process the mermaid diagram and embed as image
                mermaid_code = '\n'.join(mermaid_lines)
                
                # Try to extract diagram title
                title = "Flow Diagram"
                for ml in mermaid_lines:
                    ml_clean = ml.strip()
                    if ml_clean.startswith('graph') or ml_clean.startswith('flowchart'):
                        title = "Flow Chart"
                        break
                    elif ml_clean.startswith('erDiagram'):
                        title = "Entity Relationship Diagram"
                        break
                    elif ml_clean.startswith('sequenceDiagram'):
                        title = "Sequence Diagram"
                        break
                
                # Add a caption/title for the diagram
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f'📊 Diagram {mermaid_index}: {title}')
                run.bold = True
                run.font.color.rgb = RGBColor(0x1A, 0x23, 0x7E)
                run.font.size = Pt(11)
                
                # Download and embed the image
                img_path = get_mermaid_image_path(mermaid_code, mermaid_index)
                if img_path and os.path.exists(img_path):
                    try:
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.paragraph_format.space_after = Pt(12)
                        
                        # Add image to paragraph directly
                        # First run should contain the picture
                        run_img = p_img.add_run()
                        run_img.add_picture(img_path, width=Inches(5.5))
                    except Exception as ex:
                        p_err = doc.add_paragraph()
                        p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        run_err = p_err.add_run(f'[Error rendering diagram {mermaid_index}: {str(ex)}]')
                        run_err.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
                        run_err.italic = True
                else:
                    p_err = doc.add_paragraph()
                    p_err.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_err = p_err.add_run(f'[Diagram {mermaid_index} could not be downloaded]')
                    run_err.font.color.rgb = RGBColor(0xD3, 0x2F, 0x2F)
                    run_err.italic = True
                
                in_mermaid = False
                mermaid_lines = []
                mermaid_index += 1
                i += 1
                continue
            mermaid_lines.append(line)
            i += 1
            continue
        
        if in_code_block:
            if line.strip() == '```':
                # Render code block
                for cl in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.left_indent = Cm(0.8)
                    run = p.add_run(cl)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
                
                # Add small gap after code block
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                
                in_code_block = False
                code_lines = []
                i += 1
                continue
            code_lines.append(line)
            i += 1
            continue
        
        stripped = line.strip()
        
        # ── Skip empty lines ──
        if not stripped:
            i += 1
            continue
        
        # ── Horizontal rule ──
        if stripped == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line
            run = p.add_run('─' * 80)
            run.font.size = Pt(6)
            run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
            i += 1
            continue
        
        # ── Headings ──
        if stripped.startswith('#'):
            match = re.match(r'^(#{1,4})\s+(.*)', stripped)
            if match:
                level = len(match.group(1))
                heading_text = match.group(2)
                # Remove emoji at start for cleaner heading but keep it
                p = doc.add_heading('', level=level)
                add_formatted_text(p, heading_text)
                i += 1
                continue
        
        # ── Blockquotes / Alerts ──
        if stripped.startswith('>'):
            alert_text = stripped.lstrip('> ').strip()
            
            # Check for GitHub-style alerts
            alert_type = None
            alert_color = RGBColor(0x42, 0x42, 0x42)
            bg_color = "E3F2FD"
            
            if '[!NOTE]' in alert_text:
                alert_type = '📝 NOTE'
                alert_color = RGBColor(0x0D, 0x47, 0xA1)
                bg_color = "E3F2FD"
                i += 1
                # Get the content of the alert (next > lines)
                alert_content_parts = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    alert_content_parts.append(lines[i].strip().lstrip('> ').strip())
                    i += 1
                alert_content = ' '.join(alert_content_parts)
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(f'{alert_type}: ')
                run.bold = True
                run.font.color.rgb = alert_color
                run.font.size = Pt(10)
                add_formatted_text(p, alert_content)
                continue
            
            elif '[!IMPORTANT]' in alert_text:
                alert_type = '⚠️ IMPORTANT'
                alert_color = RGBColor(0xE6, 0x51, 0x00)
                i += 1
                alert_content_parts = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    alert_content_parts.append(lines[i].strip().lstrip('> ').strip())
                    i += 1
                alert_content = ' '.join(alert_content_parts)
                
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(f'{alert_type}: ')
                run.bold = True
                run.font.color.rgb = alert_color
                run.font.size = Pt(10)
                add_formatted_text(p, alert_content)
                continue
            
            else:
                # Regular blockquote
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                quote_text = alert_text
                add_formatted_text(p, quote_text)
                p.runs[0].italic = True if p.runs else None
                i += 1
                continue
        
        # ── Tables ──
        if stripped.startswith('|') and stripped.endswith('|'):
            table_lines_collected = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines_collected.append(lines[i])
                i += 1
            
            rows = parse_table(table_lines_collected)
            if not rows:
                continue
            
            num_cols = len(rows[0])
            # Ensure all rows have same number of columns
            for r_idx in range(len(rows)):
                while len(rows[r_idx]) < num_cols:
                    rows[r_idx].append('')
            
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = table.rows[r_idx].cells[c_idx]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                    
                    # Clean cell text
                    clean_text = cell_text.strip()
                    
                    if r_idx == 0:
                        # Header row styling
                        set_cell_shading(cell, "1A237E")
                        run = p.add_run(clean_text)
                        run.bold = True
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.name = 'Calibri'
                    else:
                        # Alternate row colors
                        if r_idx % 2 == 0:
                            set_cell_shading(cell, "F5F5F5")
                        add_formatted_text(p, clean_text)
                        for run in p.runs:
                            run.font.size = Pt(9)
                            run.font.name = 'Calibri'
            
            # Add small gap after table
            doc.add_paragraph().paragraph_format.space_after = Pt(4)
            continue
        
        # ── Unordered list items ──
        if re.match(r'^[\-\*]\s', stripped):
            list_text = stripped[2:].strip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.5)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run('• ')
            run.font.size = Pt(11)
            add_formatted_text(p, list_text)
            i += 1
            continue
        
        # ── Ordered list items ──
        if re.match(r'^\d+\.\s', stripped):
            match = re.match(r'^(\d+)\.\s(.*)', stripped)
            if match:
                num = match.group(1)
                list_text = match.group(2)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1.2)
                p.paragraph_format.first_line_indent = Cm(-0.5)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(f'{num}. ')
                run.bold = True
                run.font.size = Pt(11)
                add_formatted_text(p, list_text)
                i += 1
                continue
        
        # ── Regular paragraph ──
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1
    
    # Save
    try:
        doc.save(OUTPUT_FILE)
        print(f"✅ DOCX file saved to: {OUTPUT_FILE}")
        print(f"   File size: {os.path.getsize(OUTPUT_FILE) / 1024:.1f} KB")
    except PermissionError:
        alternative_path = OUTPUT_FILE.replace(".docx", "_v2.docx")
        print(f"⚠️ Warning: Permission denied when saving to '{OUTPUT_FILE}'.")
        print(f"   The file might be open in another application (like Word).")
        print(f"   Saving to alternative path instead: {alternative_path}")
        doc.save(alternative_path)
        print(f"✅ DOCX file saved to: {alternative_path}")
        print(f"   File size: {os.path.getsize(alternative_path) / 1024:.1f} KB")


if __name__ == '__main__':
    with open(INPUT_FILE, 'r', encoding='utf-8') as f:
        content = f.read()
    
    print("📄 Converting Smart Classroom Documentation to DOCX...")
    print(f"   Input: {INPUT_FILE}")
    print(f"   Output: {OUTPUT_FILE}")
    create_docx(content)
    print("🎉 Done!")
