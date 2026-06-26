"""
Smart Classroom LMS — Complete Project Documentation Generator
===============================================================
Generates a comprehensive Word document (.docx) with:
  - Cover page
  - Table of Contents
  - 14 main sections with full content
  - 21+ embedded diagrams
  - Development history appendix
  
Author: Yaswanth Munagoti
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── Configuration ───────────────────────────────────────────────────────────
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
IMAGES_DIR = os.path.join(BASE_DIR, 'images')
OUTPUT_FILE = os.path.join(BASE_DIR, 'Smart_Classroom_Complete_Documentation.docx')

# ─── Color Palette ───────────────────────────────────────────────────────────
COLOR_NAVY      = RGBColor(0x1A, 0x23, 0x7E)
COLOR_BLUE      = RGBColor(0x0D, 0x47, 0xA1)
COLOR_GREEN     = RGBColor(0x1B, 0x5E, 0x20)
COLOR_PURPLE    = RGBColor(0x4A, 0x14, 0x8C)
COLOR_ORANGE    = RGBColor(0xE6, 0x51, 0x00)
COLOR_RED       = RGBColor(0xC6, 0x28, 0x28)
COLOR_DARK      = RGBColor(0x21, 0x21, 0x21)
COLOR_GRAY      = RGBColor(0x61, 0x61, 0x61)
COLOR_LIGHT     = RGBColor(0x75, 0x75, 0x75)
COLOR_WHITE     = RGBColor(0xFF, 0xFF, 0xFF)

HEADER_BG       = "1A237E"
ALT_ROW_BG      = "F5F5F5"
LIGHT_BLUE_BG   = "E3F2FD"
LIGHT_GREEN_BG  = "E8F5E9"
LIGHT_ORANGE_BG = "FFF3E0"
LIGHT_RED_BG    = "FFEBEE"
LIGHT_PURPLE_BG = "F3E5F5"


# ═══════════════════════════════════════════════════════════════════════════════
#  HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def setup_styles(doc):
    """Configure document styles."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.15

    configs = [
        (1, 24, COLOR_NAVY),
        (2, 18, COLOR_BLUE),
        (3, 14, COLOR_GREEN),
        (4, 12, COLOR_PURPLE),
    ]
    for level, size, color in configs:
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Calibri'
        hs.font.bold = True
        hs.font.size = Pt(size)
        hs.font.color.rgb = color
        hs.paragraph_format.space_before = Pt(18 if level <= 2 else 12)
        hs.paragraph_format.space_after = Pt(6)


def add_paragraph(doc, text, bold=False, italic=False, size=11, color=None,
                  alignment=None, space_before=0, space_after=4, indent=0):
    """Add a styled paragraph."""
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text, indent=1.2):
    """Add a bullet point."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run('•  ')
    run.font.size = Pt(11)
    _add_rich_text(p, text)
    return p


def add_numbered(doc, num, text, indent=1.2):
    """Add a numbered list item."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.4)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(f'{num}.  ')
    run.bold = True
    run.font.size = Pt(11)
    _add_rich_text(p, text)
    return p


def _add_rich_text(paragraph, text):
    """Parse **bold** and `code` in text and add runs."""
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = COLOR_GREEN
        else:
            run = paragraph.add_run(part)
            run.font.size = Pt(11)


def add_code_block(doc, code_text):
    """Add a formatted code block."""
    for line in code_text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_DARK
    doc.add_paragraph()  # gap after code


def add_table(doc, headers, rows, col_widths=None):
    """Add a styled table with header and alternating rows."""
    num_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for c_idx, header in enumerate(headers):
        cell = table.rows[0].cells[c_idx]
        cell.text = ''
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        set_cell_shading(cell, HEADER_BG)
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = COLOR_WHITE
        run.font.name = 'Calibri'

    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= num_cols:
                break
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'

    # Set column widths if provided
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(width)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, filename, width_inches=5.5, caption=None):
    """Add a centered image with optional caption."""
    img_path = os.path.join(IMAGES_DIR, filename)
    if not os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'[Image not found: {filename}]')
        run.font.color.rgb = COLOR_RED
        run.italic = True
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(img_path, width=Inches(width_inches))

    if caption:
        pc = doc.add_paragraph()
        pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pc.paragraph_format.space_after = Pt(10)
        run_c = pc.add_run(caption)
        run_c.italic = True
        run_c.font.size = Pt(9)
        run_c.font.color.rgb = COLOR_LIGHT


def add_note(doc, text, note_type='NOTE'):
    """Add a styled note/alert box."""
    colors = {
        'NOTE': (COLOR_BLUE, LIGHT_BLUE_BG),
        'IMPORTANT': (COLOR_ORANGE, LIGHT_ORANGE_BG),
        'TIP': (COLOR_GREEN, LIGHT_GREEN_BG),
        'WARNING': (COLOR_RED, LIGHT_RED_BG),
    }
    color, _ = colors.get(note_type, (COLOR_BLUE, LIGHT_BLUE_BG))
    icons = {'NOTE': '📝', 'IMPORTANT': '⚠️', 'TIP': '💡', 'WARNING': '🔴'}
    icon = icons.get(note_type, '📝')

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f'{icon} {note_type}: ')
    run.bold = True
    run.font.color.rgb = color
    run.font.size = Pt(10)
    _add_rich_text(p, text)


def add_divider(doc):
    """Add a horizontal divider line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run('─' * 85)
    run.font.size = Pt(6)
    run.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)


def add_page_break(doc):
    """Add a page break."""
    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def build_cover_page(doc):
    """Build a professional cover page."""
    # Add several blank lines to push content down
    for _ in range(6):
        doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Project emoji/icon line
    add_paragraph(doc, '🎓', size=36,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('Smart Classroom')
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = COLOR_NAVY
    run.font.name = 'Calibri'

    # Subtitle
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(24)
    run2 = p2.add_run('Learning Management System')
    run2.font.size = Pt(20)
    run2.font.color.rgb = COLOR_BLUE
    run2.font.name = 'Calibri'

    # Divider
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(24)
    run3 = p3.add_run('━' * 40)
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)

    # Document type
    add_paragraph(doc, 'Complete Project Documentation',
                  size=16, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

    # Author
    add_paragraph(doc, 'Prepared by',
                  size=12, color=COLOR_LIGHT, italic=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)

    p_author = doc.add_paragraph()
    p_author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_author.paragraph_format.space_after = Pt(30)
    run_a = p_author.add_run('Yaswanth Munagoti')
    run_a.bold = True
    run_a.font.size = Pt(18)
    run_a.font.color.rgb = COLOR_NAVY
    run_a.font.name = 'Calibri'

    # Tech stack line
    add_paragraph(doc, 'Django  ·  Python  ·  SQLite  ·  Bootstrap 5  ·  Gemini AI',
                  size=11, color=COLOR_LIGHT,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)

    # Date
    add_paragraph(doc, 'June 2026',
                  size=12, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════════════════

def build_toc(doc):
    """Insert a Word-native Table of Contents field."""
    doc.add_heading('Table of Contents', level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)

    # Insert TOC field code
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    fld_char_separate = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run1 = p.add_run()
    run1._r.append(fld_char_begin)
    run2 = p.add_run()
    run2._r.append(instr_text)
    run3 = p.add_run()
    run3._r.append(fld_char_separate)

    # Placeholder text
    run_placeholder = p.add_run('Right-click and select "Update Field" to generate Table of Contents')
    run_placeholder.font.color.rgb = COLOR_GRAY
    run_placeholder.italic = True

    run4 = p.add_run()
    run4._r.append(fld_char_end)

    add_note(doc, 'To generate the Table of Contents in Word: Right-click the text above → select **"Update Field"** → choose **"Update entire table"**.', 'TIP')

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 1: PROJECT OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_1(doc):
    doc.add_heading('1. Project Overview', level=1)

    add_paragraph(doc,
        'Smart Classroom is a comprehensive web-based Learning Management System (LMS) '
        'built with Django that connects administrators, teachers, and students in a '
        'digital classroom environment. The platform provides end-to-end academic management '
        'including course creation, assignments, quizzes, online coding tests, attendance tracking, '
        'AI-powered tutoring, real-time chat, certificates, and analytics — all within a single, '
        'unified web application.')

    add_paragraph(doc,
        'The system is designed to replicate and enhance the traditional classroom experience '
        'by providing digital tools for every aspect of teaching and learning. It supports '
        '3 user roles (Admin, Teacher, Student), 3 approval workflows, and features an '
        'anti-cheat exam system for secure online assessments.')

    doc.add_heading('Key Highlights', level=3)
    highlights = [
        '**12 Django Applications** working together as a cohesive system',
        '**24 Database Models** with 29 foreign key and 3 many-to-many relationships',
        '**105+ URL Endpoints** covering all CRUD operations and API endpoints',
        '**50+ HTML Templates** with responsive Bootstrap 5 design',
        '**3 User Roles** (Admin, Teacher, Student) with granular permissions',
        '**3 Approval Workflows** for user registration, course creation, and enrollment',
        '**AI Integration** via Google Gemini for per-course AI tutoring',
        '**Anti-Cheat Exam System** with fullscreen lock, tab detection, and auto-submit',
        '**4 Assessment Types**: Assignments, Quizzes, Coding Tests, Projects',
        '**Real-time Chat** with direct messages and per-course group chat',
        '**Data Export** in CSV, Excel, and PDF formats',
        '**Light/Dark Mode** with CSS variable-based theme toggle',
    ]
    for h in highlights:
        add_bullet(doc, h)

    doc.add_heading('Feature Summary', level=3)
    add_table(doc,
        ['#', 'Feature', 'Description'],
        [
            ['1',  'User Management',            '3 roles (Admin, Teacher, Student) with approval workflows'],
            ['2',  'Department & Course Mgmt',    'Departments → Courses with unique codes, approval system'],
            ['3',  'Enrollment System',           'Students request enrollment → Teacher/Admin approves'],
            ['4',  'Assignments',                 'Create, submit (file/GitHub), grade with marks & feedback'],
            ['5',  'Quizzes',                     'MCQ-based quizzes with auto-scoring, time limits, date windows'],
            ['6',  'Online Coding Tests',         'In-browser coding IDE with sample/hidden test cases, auto-evaluation'],
            ['7',  'Attendance',                  'Teacher marks or student self-marks (with deadlines)'],
            ['8',  'Study Materials',             'Upload PDF, DOC, PPT, ZIP, PY files for download'],
            ['9',  'Recorded Classes',            'Video uploads (file or URL) with streaming player'],
            ['10', 'Team Projects',               'Project submissions with GitHub links, ZIP uploads, grading'],
            ['11', 'Chat System',                 '1-on-1 direct messages + per-course group chat (AJAX real-time)'],
            ['12', 'AI Tutor',                    'Per-course AI chatbot powered by Google Gemini'],
            ['13', 'Virtual Sessions',            'Schedule online meeting links with auto-attendance'],
            ['14', 'Certificates',                'Issue badges (Completion, Excellence, Topper, Participation)'],
            ['15', 'Reports & Leaderboard',       'Report cards, grade exports (CSV/Excel), rankings'],
            ['16', 'Notification System',         'In-app notifications for all key events'],
            ['17', 'Admin Panel',                 'Custom admin dashboard for managing everything'],
            ['18', 'Light/Dark Mode',             'CSS variable-based theme toggle'],
        ],
        col_widths=[1.2, 4, 10.8]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 2: TECHNOLOGY STACK
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_2(doc):
    doc.add_heading('2. Technology Stack & Requirements', level=1)

    add_paragraph(doc,
        'The Smart Classroom LMS is built using a modern Python web development stack. '
        'Django serves as the core web framework providing the MVC architecture, ORM, '
        'authentication, and templating engine. The frontend uses Bootstrap 5 for responsive '
        'design with custom CSS for theming.')

    doc.add_heading('Software Dependencies', level=3)
    add_table(doc,
        ['Technology', 'Purpose', 'Version'],
        [
            ['Python',               'Backend programming language',              '3.10+'],
            ['Django',               'Web framework (MVT architecture)',           '≥4.2, <5.0'],
            ['SQLite3',              'Database (built into Django/Python)',         'Built-in'],
            ['Pillow',               'Image processing (profile pics, banners)',   '≥9.0'],
            ['WhiteNoise',           'Static file serving in production',          '≥6.0'],
            ['Gunicorn',             'Production WSGI server',                     '≥20.0'],
            ['WeasyPrint',           'PDF generation for certificates/reports',    '≥59.0'],
            ['openpyxl',             'Excel export for marks/reports',             '≥3.1'],
            ['Google Generative AI', 'Gemini AI integration (AI Tutor)',           'Latest'],
            ['Bootstrap 5',         'Frontend CSS framework (via CDN)',            '5.x'],
            ['Chart.js',            'Dashboard analytics charts',                  'Latest'],
            ['Font Awesome',        'Icon library for UI elements',                '6.x'],
        ],
        col_widths=[4, 7, 3]
    )

    doc.add_heading('requirements.txt', level=3)
    add_code_block(doc, """Django>=4.2,<5.0
Pillow>=9.0
whitenoise>=6.0
gunicorn>=20.0
WeasyPrint>=59.0
openpyxl>=3.1""")

    doc.add_heading('Key Configuration (settings.py)', level=3)
    add_table(doc,
        ['Setting', 'Value'],
        [
            ['Database',        'SQLite3 (db.sqlite3)'],
            ['User Model',      'users.CustomUser (extends AbstractUser)'],
            ['Time Zone',       'Asia/Kolkata'],
            ['Login Redirect',  '/users/dashboard/'],
            ['Logout Redirect', '/users/login/'],
            ['Media Root',      'BASE_DIR / "media"'],
            ['Static Root',     'BASE_DIR / "staticfiles"'],
            ['CSRF Origins',    'localhost, ngrok tunnels'],
        ],
        col_widths=[5, 11]
    )

    doc.add_heading('Django Applications (12 Apps)', level=3)
    add_paragraph(doc, 'The project is organized into 12 Django applications, each responsible for a specific domain:')
    add_table(doc,
        ['App Name', 'Purpose', 'Key Models'],
        [
            ['users',           'Authentication, profiles, admin panel, notifications', 'CustomUser, Notification'],
            ['courses',         'Departments, courses, enrollment, virtual sessions, AI tutor', 'Department, Course, Enrollment, VirtualSession, AITutorMessage'],
            ['assignments',     'Assignment creation, submission, grading', 'Assignment, Submission'],
            ['quiz',            'MCQ quizzes with auto-scoring', 'Quiz, Question, Choice, QuizAttempt, StudentAnswer'],
            ['tests',           'Online coding tests with test cases', 'Test, CodingQuestion, TestCase, StudentResponse'],
            ['attendance',      'Attendance sessions and records', 'AttendanceSession, AttendanceRecord'],
            ['materials',       'Study material file uploads', 'Material'],
            ['recorded_classes','Video recordings for courses', 'RecordedClass'],
            ['projects',        'Team project submissions', 'Project'],
            ['chat',            'Direct messages and group chat', 'DirectMessage, GroupMessage'],
            ['certificates',    'Digital badges and certificates', 'Certificate'],
            ['reports',         'Report cards and grade analytics', 'ReportCard'],
        ],
        col_widths=[3, 6, 7]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 3: SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_3(doc):
    doc.add_heading('3. System Architecture', level=1)

    add_paragraph(doc,
        'The Smart Classroom LMS follows Django\'s Model-View-Template (MVT) architectural '
        'pattern. The system is designed as a monolithic Django application with 12 tightly '
        'integrated apps. Each app encapsulates its own models, views, URLs, and templates, '
        'while sharing the common user authentication and course enrollment infrastructure.')

    doc.add_heading('Architecture Diagram', level=3)
    add_image(doc, 'diagram_1.png', 5.5, 'Figure 1: Smart Classroom System Architecture')

    doc.add_heading('Request Lifecycle', level=3)
    add_paragraph(doc,
        'Every HTTP request follows a well-defined pipeline through the Django framework:')

    add_code_block(doc, """Browser → HTTP Request → Django URL Router → View Function
                                                  ↓
                                           Permission Check (Role + Enrollment + Assignment)
                                                  ↓
                                           Database Query (ORM)
                                                  ↓
                                           [Optional] Gemini AI Call
                                                  ↓
                                           Template Rendering → HTML Response → Browser""")

    add_paragraph(doc,
        'The middleware stack includes security middleware, session management, CSRF protection, '
        'authentication middleware, and a custom **AccountStatusMiddleware** that intercepts '
        'requests from unapproved users and redirects them to an approval status page.')

    doc.add_heading('Project Directory Structure', level=3)
    add_code_block(doc, """smart_classroom/                    # Project root
├── smart_classroom/               # Django project settings
│   ├── settings.py               # Configuration (DB, apps, middleware)
│   ├── urls.py                   # Root URL configuration
│   └── wsgi.py                   # WSGI entry point
├── users/                         # User authentication & admin panel
├── courses/                       # Department, course, enrollment, AI tutor
├── assignments/                   # Assignment CRUD & grading
├── quiz/                          # MCQ quiz system
├── tests/                         # Online coding test engine
├── attendance/                    # Attendance tracking
├── materials/                     # Study material uploads
├── recorded_classes/              # Video recording management
├── projects/                      # Team project submissions
├── chat/                          # Direct & group messaging
├── certificates/                  # Digital badges & certificates
├── reports/                       # Report cards & analytics
├── templates/                     # Shared base templates
├── static/                        # CSS, JS, images
├── media/                         # User-uploaded files
├── images/                        # Documentation diagrams
├── db.sqlite3                     # SQLite database
├── manage.py                      # Django management script
├── requirements.txt               # Python dependencies
├── setup.py                       # Automated setup script
└── test_suite.py                  # Comprehensive test suite""")

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 4: USER ROLES & PERMISSIONS
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_4(doc):
    doc.add_heading('4. User Roles & Permissions', level=1)

    add_paragraph(doc,
        'The system implements a 3-tier role-based access control (RBAC) model. Each user '
        'is assigned exactly one role at registration, which determines their dashboard, '
        'navigation menu, and permitted actions throughout the platform.')

    doc.add_heading('Role Descriptions', level=3)

    add_paragraph(doc, '**👑 Admin** — The superuser with full system control. Admins manage departments, '
        'approve teacher/student registrations, approve courses, and have access to the Django admin panel. '
        'Admin accounts are auto-approved upon creation.', space_after=8)

    add_paragraph(doc, '**🧑‍🏫 Teacher** — Course instructors who create and manage academic content. '
        'Teachers can create courses (subject to admin approval), manage assignments, quizzes, coding tests, '
        'attendance, materials, recorded classes, and grade student submissions. Teacher accounts require '
        'admin approval before activation.', space_after=8)

    add_paragraph(doc, '**🎓 Student** — Learners who enroll in courses and complete academic activities. '
        'Students can browse departments, request enrollment, submit assignments, take quizzes/tests, '
        'view materials, chat with peers/teachers, and access AI tutoring. Student accounts require '
        'admin approval before activation.', space_after=8)

    doc.add_heading('Role Comparison Matrix', level=3)
    add_table(doc,
        ['Feature', 'Admin', 'Teacher', 'Student'],
        [
            ['Account Status',    'Auto-active',             'Needs approval',            'Needs approval'],
            ['Dashboard',         'Admin panel',             'Course overview + charts',  'Learning overview + charts'],
            ['Departments',       'Create / Edit / List',    'View only',                 'View only'],
            ['Courses',           'Create / Edit / Approve', 'Create (needs approval)',   'View & enroll'],
            ['Enrollment',        'Approve / Reject',        'Approve / Reject (own)',    'Request enrollment'],
            ['Assignments',       '—',                       'Create / Edit / Grade',     'View / Submit'],
            ['Quizzes',           '—',                       'Create / Edit / Add Qs',    'Take quiz (once)'],
            ['Online Tests',      '—',                       'Create / Edit / Add Qs',    'Take test (once)'],
            ['Attendance',        '—',                       'Create session / Mark',     'Self-mark (if allowed)'],
            ['Materials',         '—',                       'Upload / Delete',           'Download'],
            ['Recorded Classes',  '—',                       'Upload / Edit / Delete',    'View / Watch'],
            ['Projects',          '—',                       'Grade',                     'Submit / Delete own'],
            ['Chat (DM)',         'Chat with anyone',        'Chat with anyone',          'Chat with teachers & peers'],
            ['Group Chat',        '—',                       'Course group chat',         'Course group chat'],
            ['AI Tutor',          '—',                       'Access per course',         'Access per course'],
            ['Certificates',      '—',                       'Issue to students',         'View / Download own'],
            ['Reports',           '—',                       'View all / Edit / Export',  'View own report card'],
            ['Leaderboard',       'View',                    'View',                      'View'],
            ['Notifications',     'Receive',                 'Receive',                   'Receive'],
            ['Admin Panel',       '✅ Full access',           '❌',                        '❌'],
        ],
        col_widths=[3.5, 3.5, 4, 4]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 5: APPROVAL WORKFLOWS
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_5(doc):
    doc.add_heading('5. Approval Workflows', level=1)

    add_paragraph(doc,
        'The system implements 3 approval workflows to maintain quality control and ensure '
        'only authorized users access the platform. Each workflow follows a request → review → '
        'approve/reject pattern with in-app notifications at each stage.')

    doc.add_heading('Workflow 1: User Registration Approval', level=3)
    add_paragraph(doc,
        'When a new teacher or student registers, their account is created with `approval_status = pending`. '
        'They are redirected to an approval status page and cannot access the dashboard until an Admin '
        'approves their account. The custom `AccountStatusMiddleware` enforces this restriction on every request.')
    add_image(doc, 'diagram_2.png', 5.0,
              'Figure 2: User Registration Approval Workflow')

    doc.add_heading('Workflow 2: Course Creation Approval', level=3)
    add_paragraph(doc,
        'When a teacher creates a new course, it is saved with `approval_status = pending`. '
        'The course is invisible to students until an Admin approves it. The system auto-generates '
        'a unique course code (e.g., CS-ML2026) and batch identifier. Admin-created courses are auto-approved.')
    add_image(doc, 'diagram_3.png', 5.0,
              'Figure 3: Course Creation Approval Workflow')

    doc.add_heading('Workflow 3: Student Enrollment Approval', level=3)
    add_paragraph(doc,
        'When a student requests enrollment in a course, an `Enrollment` record is created with '
        '`status = pending`. The course teacher receives a notification and can approve or reject '
        'the request. Approved students get full access to all course content. The system supports '
        'soft-delete for un-enrollment (preserving student data) and seamless re-enrollment.')
    add_image(doc, 'diagram_4.png', 5.0,
              'Figure 4: Student Enrollment Approval Workflow')

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 6: DATABASE SCHEMA
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_6(doc):
    doc.add_heading('6. Complete Database Schema', level=1)

    add_paragraph(doc,
        'The Smart Classroom database consists of 24 models distributed across 12 Django apps. '
        'The schema is designed around two central entities: **CustomUser** (referenced by 17 FKs '
        'and 3 M2M relationships) and **Course** (referenced by 15 FKs). Together, these form '
        'the backbone connecting all modules of the LMS.')

    # ── 6.1 Complete ERD ──
    doc.add_heading('6.1 Complete Database Schema (ERD)', level=2)
    add_paragraph(doc,
        'The following diagram shows the complete entity-relationship diagram for all 24 models '
        'and 29 relationships in the system:')
    add_image(doc, 'complete_database_schema.png', 6.0,
              'Figure 5: Complete Database Schema — Entity Relationship Diagram')

    # ── 6.2 Module-by-module ──
    doc.add_heading('6.2 Entity Relationship Diagrams (By Module)', level=2)
    add_paragraph(doc,
        'To ensure readability, the 24-model schema has been partitioned into 7 logical modules:')

    modules = [
        ('Module 1: User Profiles & Direct Messages',
         'Manages registration status, authentication roles, in-app notifications, and direct messaging between users.',
         'diagram_5.png', 6),
        ('Module 2: Course & Enrollment Structure',
         'Handles departments, courses, student enrollment requests, virtual live sessions, per-course group chats, and AI Tutor conversation logs.',
         'diagram_6.png', 7),
        ('Module 3: Study Materials & Recorded Classes',
         'Stores uploaded files (PDFs, PPTs, code files) and recorded video links/files for course materials.',
         'diagram_7.png', 8),
        ('Module 4: Assignments & Project Management',
         'Organizes standard assignments (with grade + feedback) and capstone team project submissions.',
         'diagram_8.png', 9),
        ('Module 5: Quiz Assessment System',
         'Represents MCQ-based quizzes, questions, options (choices), attempts, and student answers.',
         'diagram_9.png', 10),
        ('Module 6: Online Coding Tests',
         'Details the coding exam engine, including questions, test cases (sample/hidden), and execution evaluation scores.',
         'diagram_10.png', 11),
        ('Module 7: Attendance, Grading & Certificates',
         'Manages student attendance tracking, issued completion/excellence badges, and aggregated report cards.',
         'diagram_11.png', 12),
    ]

    for idx, (title, desc, img, fig_num) in enumerate(modules):
        doc.add_heading(title, level=3)
        add_paragraph(doc, desc)
        add_image(doc, img, 5.5, f'Figure {fig_num}: {title} — ERD')

    # ── 6.3 Unique Constraints ──
    doc.add_heading('6.3 Unique Constraints', level=2)
    add_paragraph(doc,
        'The following unique constraints enforce data integrity across the database, '
        'preventing duplicate records in critical relationships:')
    add_table(doc,
        ['Model', 'Constraint', 'Purpose'],
        [
            ['Enrollment',       '(student, course)',         'One enrollment per student per course'],
            ['Submission',       '(assignment, student)',     'One submission per assignment'],
            ['AttendanceSession','(course, date)',            'One session per day per course'],
            ['AttendanceRecord', '(session, student)',        'One record per student per session'],
            ['QuizAttempt',      '(quiz, student)',           'One attempt per student per quiz'],
            ['StudentAnswer',    '(attempt, question)',       'One answer per question'],
            ['StudentResponse',  '(student, test)',           'One response per test'],
            ['Certificate',      '(student, course, badge_type)', 'One badge type per student per course'],
            ['ReportCard',       '(student, course)',         'One report per student per course'],
        ],
        col_widths=[3.5, 4.5, 7]
    )

    # ── 6.4 Summary Statistics ──
    doc.add_heading('6.4 Summary Statistics', level=2)
    add_table(doc,
        ['Metric', 'Count'],
        [
            ['Total Django Apps',              '12'],
            ['Total Database Models',          '24'],
            ['Foreign Key Relationships',      '29'],
            ['Many-to-Many Relationships',     '3'],
            ['Central Model',                  'Course (referenced by 15 FKs)'],
            ['User Model',                     'CustomUser (referenced by 17 FKs + 3 M2M)'],
        ],
        col_widths=[7, 9]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 7: APPLICATION FLOWS
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_7(doc):
    doc.add_heading('7. Application Flows — Screen by Screen', level=1)

    add_paragraph(doc,
        'This section documents the complete user journey through every major feature of the platform. '
        'Each flow diagram shows the step-by-step interaction between the user and the system.')

    # Flow 1
    doc.add_heading('Flow 1: Registration → Approval → Dashboard', level=3)
    add_paragraph(doc,
        'New users register by selecting their role (Teacher/Student), filling in profile details, '
        'and submitting the form. The system creates an account with pending status. Once an Admin '
        'approves, the user is redirected to their role-specific dashboard on next login.')
    add_image(doc, 'diagram_12.png', 5.0,
              'Figure 13: Registration → Approval → Dashboard Flow')

    # Flow 2 - Teacher Dashboard
    doc.add_heading('Flow 2: Teacher Dashboard — Data Overview', level=3)
    add_paragraph(doc, 'The teacher dashboard provides a comprehensive overview of all teaching activities:')
    add_table(doc,
        ['Dashboard Section', 'Data Displayed'],
        [
            ['My Courses',        'All courses taught with student/assignment/quiz/pending counts'],
            ['Total Students',    'Sum of enrolled students across all courses'],
            ['Total Assignments', 'Count of all assignments created'],
            ['Total Quizzes',     'Count of all quizzes created'],
            ['Pending vs Graded', 'Submission breakdown (assignments + projects)'],
            ['Attendance Rate',   'Overall present/absent percentages'],
            ['Course Status',     'Active / Inactive / Results Released counts'],
            ['Pending Reviews',   'Top 5 ungraded submissions requiring action'],
            ['Charts',            'Student counts, assignment counts, submission breakdown, attendance (JSON-powered Chart.js)'],
        ],
        col_widths=[4, 12]
    )

    # Flow 3 - Student Dashboard
    doc.add_heading('Flow 3: Student Dashboard — Data Overview', level=3)
    add_paragraph(doc, 'The student dashboard focuses on learning progress and upcoming tasks:')
    add_table(doc,
        ['Dashboard Section', 'Data Displayed'],
        [
            ['Enrolled Courses',     'Courses with approved enrollment'],
            ['Pending Work',         'Unsubmitted assignments + unattempted quizzes + tests + projects'],
            ['Overall Performance',  'Average % across all graded work'],
            ['Attendance %',         'Overall attendance rate across all courses'],
            ['Student Rank',         'Ranking among all students (leaderboard position)'],
            ['Upcoming Items',       'Next 4 deadlines (assignments, quizzes, tests, projects)'],
            ['Charts',               'Weekly activity (7-day), course performance, workload breakdown'],
        ],
        col_widths=[4, 12]
    )

    # Flow 4
    doc.add_heading('Flow 4: Course Creation → Enrollment', level=3)
    add_paragraph(doc,
        'Teachers create courses by filling a form with title, description, department, dates, and '
        'cover image. After admin approval, students can browse departments, find the course, and '
        'request enrollment. The teacher gets a notification and approves/rejects the request.')
    add_image(doc, 'diagram_13.png', 5.0, 'Figure 14: Course Creation → Enrollment Flow')

    # Flow 5
    doc.add_heading('Flow 5: Assignment Lifecycle', level=3)
    add_paragraph(doc,
        'Teachers create assignments with title, description, deadline, total marks, and optional '
        'file attachments. Students submit work (file upload or GitHub link) before the deadline. '
        'The teacher grades each submission with marks and feedback. Students receive notification '
        'of their grades.')
    add_image(doc, 'diagram_14.png', 5.0, 'Figure 15: Assignment Lifecycle Flow')

    # Flow 6
    doc.add_heading('Flow 6: Quiz Lifecycle', level=3)
    add_paragraph(doc,
        'Teachers create MCQ quizzes with configurable time limits, start/end dates, and question sets. '
        'Each question has multiple choices with one correct answer. Students take the quiz once within '
        'the allowed window. Scoring is automatic based on correct answers.')
    add_image(doc, 'diagram_15.png', 5.0, 'Figure 16: Quiz Lifecycle Flow')

    # Flow 7
    doc.add_heading('Flow 7: Online Coding Test (Unique Feature)', level=3)
    add_paragraph(doc,
        'This is a standout feature of the LMS — an in-browser coding IDE where students write and '
        'submit code solutions. Teachers define coding questions with sample test cases (visible to students) '
        'and hidden test cases (for evaluation). The system runs student code against all test cases '
        'server-side and assigns a weighted score based on passed test cases.')
    add_image(doc, 'diagram_16.png', 5.0, 'Figure 17: Online Coding Test Flow')

    # Flow 8
    doc.add_heading('Flow 8: Attendance Flow', level=3)
    add_paragraph(doc,
        'Teachers create attendance sessions for specific dates. They can mark attendance manually '
        'or enable student self-marking within a configurable time window. Attendance records are '
        'tracked per student per session with unique constraints preventing duplicates.')
    add_image(doc, 'diagram_17.png', 5.0, 'Figure 18: Attendance Tracking Flow')

    # Flow 9
    doc.add_heading('Flow 9: Chat System', level=3)
    add_paragraph(doc,
        'The chat system supports two modes: direct messages (1-on-1) and course group chat. '
        'Messages are sent and received via AJAX polling for real-time feel. The system includes '
        'emoji picker, file upload support, and proper access control.')
    add_image(doc, 'diagram_18.png', 5.0, 'Figure 19: Chat System Flow')
    add_note(doc, '**Access Control**: Students can only DM teachers of their enrolled courses, '
             'peer students in the same courses, and admins. Teachers and admins can DM anyone.', 'NOTE')

    # Flow 10
    doc.add_heading('Flow 10: Certificate Issuance', level=3)
    add_paragraph(doc,
        'At the end of a course, teachers can issue digital badges/certificates to students. '
        'Four badge types are available: Completion, Excellence, Topper, and Participation. '
        'Each certificate is downloadable as a PDF and includes a unique verification. '
        'When a certificate is issued, a corresponding digital badge appears on the student\'s profile.')
    add_image(doc, 'diagram_19.png', 5.0, 'Figure 20: Certificate Issuance Flow')

    # Flow 11
    doc.add_heading('Flow 11: Reports & Grading', level=3)
    add_paragraph(doc,
        'Teachers can generate report cards that aggregate scores across assignments, quizzes, '
        'coding tests, and projects for each student. Reports can be exported as Excel spreadsheets '
        'with color-coded formatting and weekly data breakdown.')
    add_image(doc, 'diagram_20.png', 5.0, 'Figure 21: Reports & Grading Flow')

    # Flow 12
    doc.add_heading('Flow 12: Data Export', level=3)
    add_paragraph(doc,
        'The system supports exporting data in three formats: CSV (simple tabular data), '
        'Excel (color-coded with multiple sheets and weekly analysis), and PDF (for certificates '
        'and report cards). Teachers can download comprehensive grade reports from the course detail page.')
    add_image(doc, 'diagram_21.png', 5.0, 'Figure 22: Data Export Flow')

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 8: STEP-BY-STEP WALKTHROUGH
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_8(doc):
    doc.add_heading('8. Step-by-Step: How to Create a Course (Example)', level=1)

    add_paragraph(doc,
        'This section walks through a complete course lifecycle from creation to completion, '
        'demonstrating how the system works in practice.')

    doc.add_heading('Scenario', level=3)
    add_note(doc, '**Teacher "Dr. Sharma"** wants to create a course called '
             '**"Machine Learning Fundamentals"** under the **Computer Science** department and invite students.', 'NOTE')

    doc.add_heading('Step 1: Login as Teacher', level=3)
    add_numbered(doc, 1, 'Open browser → Go to `http://localhost:8000`')
    add_numbered(doc, 2, 'Redirected to `/users/login/`')
    add_numbered(doc, 3, 'Enter credentials: `teacher1` / `pass123`')
    add_numbered(doc, 4, 'Click **Login** → Redirected to **Teacher Dashboard**')

    doc.add_heading('Step 2: Create the Course', level=3)
    add_numbered(doc, 1, 'Click **"Create Course"** on the dashboard')
    add_numbered(doc, 2, 'Navigate to `/courses/create/`')
    add_numbered(doc, 3, 'Fill in the form:')

    add_table(doc,
        ['Field', 'Example Value'],
        [
            ['Title',         'Machine Learning Fundamentals'],
            ['Description',   'Learn the basics of ML including regression, classification, clustering, and neural networks.'],
            ['Department',    'Computer Science'],
            ['Max Students',  '50'],
            ['Start Date',    '2026-06-15'],
            ['End Date',      '2026-09-15'],
            ['Project Start', '2026-08-01'],
            ['Project End',   '2026-09-10'],
            ['Cover Image',   'Upload a banner image'],
        ],
        col_widths=[4, 12]
    )

    add_numbered(doc, 4, 'Click **Submit**')
    add_note(doc, 'The course is created with `approval_status = pending`. An Admin must approve it '
             'before students can see it. The system auto-generates a unique `course_code` (e.g., CS-ML2026) '
             'and `batch` identifier.', 'IMPORTANT')

    doc.add_heading('Step 3: Admin Approves the Course', level=3)
    add_numbered(doc, 1, 'Admin logs in → Goes to **Admin Panel** → **Courses**')
    add_numbered(doc, 2, 'Sees the pending course → Clicks **Approve**')
    add_numbered(doc, 3, 'Course is now visible to students')

    doc.add_heading('Step 4: Students Enroll', level=3)
    add_numbered(doc, 1, 'Student logs in → Browses **Departments** → Clicks **Computer Science**')
    add_numbered(doc, 2, 'Sees **"Machine Learning Fundamentals"** in the course list')
    add_numbered(doc, 3, 'Clicks the course → Sees the **About page**')
    add_numbered(doc, 4, 'Clicks **"Enroll"** → Enrollment request created (`status = pending`)')
    add_numbered(doc, 5, '**Teacher gets notification**: "Student X requested enrollment"')
    add_numbered(doc, 6, 'Teacher clicks **Approve** → Student now has full access')

    doc.add_heading('Step 5: Add Content to the Course', level=3)
    add_paragraph(doc, 'After enrollment, the teacher adds content from the Course Detail page:')
    add_table(doc,
        ['Action', 'What Teacher Does', 'What Students See'],
        [
            ['Create Assignment', 'Title: "Linear Regression", Due: June 25, Marks: 100', 'Assignment appears in dashboard "Pending"'],
            ['Upload Material',   'Upload "ML_Notes.pdf" (type: pdf)',                     'Material available for download'],
            ['Create Quiz',       'Title: "ML Basics Quiz", Duration: 30 min, Add 10 MCQs','Quiz appears in dashboard "Upcoming"'],
            ['Create Coding Test','Title: "Implement KNN", Time: 60 min, Add test cases',  'Test appears in course detail'],
            ['Upload Recording',  'Upload lecture video or paste YouTube URL',              'Recording available in course'],
            ['Schedule Session',  'Set meeting link + date/time',                           'Session listed, auto-attendance'],
            ['Open AI Tutor',     'Available immediately per course',                       'Students can chat with AI'],
        ],
        col_widths=[3.5, 6, 6]
    )

    doc.add_heading('Step 6: Student Submits Assignment', level=3)
    add_numbered(doc, 1, 'Student opens **Assignment Detail** → Clicks **Submit**')
    add_numbered(doc, 2, 'Uploads a `.py` file **OR** pastes a GitHub link')
    add_numbered(doc, 3, 'System enforces hard deadline (server-side check)')
    add_numbered(doc, 4, 'Submission saved with timestamp')

    doc.add_heading('Step 7: Teacher Grades', level=3)
    add_numbered(doc, 1, 'Teacher opens **Assignment Detail** → Sees all submissions')
    add_numbered(doc, 2, 'For each student, enters **Grade**: 85 (out of 100)')
    add_numbered(doc, 3, 'Enters **Feedback**: "Good implementation, but missing edge case handling."')
    add_numbered(doc, 4, 'Clicks **Save** → Student gets notification: "Your assignment scored 85/100"')

    doc.add_heading('Step 8: End of Course', level=3)
    add_table(doc,
        ['Action', 'Result'],
        [
            ['Generate Report Cards', 'Score summary per student (assignments + quizzes + tests)'],
            ['Export to Excel',       'Color-coded spreadsheet with all marks + weekly analysis'],
            ['Issue Certificates',    'Select badge type per student (Completion/Excellence/Topper)'],
            ['Release Results',       'Toggle results_released → students can see all scores'],
        ],
        col_widths=[5, 11]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 9: AI INTEGRATION
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_9(doc):
    doc.add_heading('9. AI Integration', level=1)

    add_paragraph(doc,
        'Smart Classroom integrates artificial intelligence in three key areas: an AI-powered '
        'tutoring chatbot with dynamic RAG context retrieval, an automated coding test evaluation engine, '
        'and a teacher assistant for AI-assisted assignment grading. Both RAG and grading features connect '
        'directly to Google\'s Gemini API using Python\'s built-in urllib libraries (zero dependencies), '
        'ensuring absolute lightweight operation and a safe local fallback to prevent server errors.')

    doc.add_heading('AI Tutor (Per-Course Chatbot with RAG)', level=3)
    add_paragraph(doc,
        'Each course has a dedicated AI Tutor powered by Google Gemini. Students and teachers '
        'can chat with the AI about course-specific topics. To make the bot specialized for each course, '
        'the system implements a custom Retrieval-Augmented Generation (RAG) engine. When a user asks a '
        'question, the RAG engine searches uploaded materials (PDFs and text files) for the course, chunks '
        'the text, ranks them by term similarity matching, and appends the top matching sections to the Gemini prompt.')

    add_table(doc,
        ['Property', 'Detail'],
        [
            ['Model',        'Google Gemini (direct API integration via urllib.request)'],
            ['URL',          '/courses/<course_pk>/ai-tutor/'],
            ['RAG Engine',   'Extracts text from PDF/code/text materials, chunks content, scores and retrieves top matches based on query terms.'],
            ['Caching',      'Parsed text documents are cached in-memory with file modification time keys to speed up subsequent queries.'],
            ['Persistence',  'Full conversation saved per user per course in the database (AITutorMessage model)'],
            ['Fallback',     'If Gemini API is blocked or key is missing, automatically runs in Local Fallback Mode showing simulated replies.'],
            ['Access',       'Only enrolled students and course teachers can access'],
        ],
        col_widths=[3, 13]
    )

    doc.add_heading('AI Assignment Grading Assistant', level=3)
    add_paragraph(doc,
        'When grading student assignment submissions, teachers can leverage the AI Grading Assistant. '
        'By clicking the "Generate AI Suggestion" button, an AJAX request triggers the server to read '
        'the student\'s submission text or code and pass it along with the assignment parameters to Gemini. '
        'The AI returns a recommended score and a structured feedback draft in Markdown, which the teacher can edit and approve.')

    add_table(doc,
        ['Property', 'Detail'],
        [
            ['Model',        'Google Gemini (MIME response type restricted to JSON)'],
            ['URL',          '/assignments/submission/<submission_pk>/ai-grade/'],
            ['API Fallback', 'Catches request blocks/network failures to return mock grading and comments based on assignment tags with zero downtime.'],
            ['Teacher Control','AI suggestions are only loaded into the scoring form inputs; nothing is saved until the teacher clicks "Save Grade".'],
        ],
        col_widths=[3, 13]
    )

    doc.add_heading('Coding Test Evaluation Engine', level=3)
    add_paragraph(doc,
        'The online coding test system includes an automated evaluation engine that runs '
        'student-submitted code against predefined test cases on the server side.')

    add_table(doc,
        ['Property', 'Detail'],
        [
            ['Function',  'evaluate_coding_question() in tests app'],
            ['Process',   'Takes student code → runs against hidden test cases → compares output'],
            ['Scoring',   'Weighted test cases, partial credit possible based on cases passed'],
            ['Security',  'Server-side execution only; code never runs in the browser'],
            ['Languages', 'Python (primary), extensible to other languages'],
        ],
        col_widths=[3, 13]
    )

    doc.add_heading('Anti-Cheat Exam System', level=3)
    add_paragraph(doc,
        'The platform implements a comprehensive anti-cheat system for quizzes and coding tests '
        'to maintain exam integrity. This was one of the most critical features developed, '
        'addressing multiple security vulnerabilities:')

    add_table(doc,
        ['Feature', 'Implementation'],
        [
            ['Fullscreen Lock',     'Exam auto-enters fullscreen on start; exit triggers warning'],
            ['Tab Switch Detection', 'visibilitychange event listener detects tab switches'],
            ['Window Blur Detection','blur event detects focus loss (half-screen, alt-tab, etc.)'],
            ['Warning Counter',      '2 warnings displayed with overlay before action'],
            ['Auto-Submit',          'On 3rd violation, form is automatically submitted with saved answers'],
            ['UI Lockdown',          'Sidebar, navbar, and footer hidden during exam (body.exam-active CSS class)'],
            ['Pre-Exam Instructions','Rules overlay shown before exam starts with "I Agree" button'],
            ['Browser X Detection',  'document.fullscreenElement === null check catches all exit methods'],
        ],
        col_widths=[4, 12]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 10: ADMIN PANEL
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_10(doc):
    doc.add_heading('10. Admin Panel', level=1)

    add_paragraph(doc,
        'The Admin Panel (`/users/admin-panel/`) provides a complete management interface for '
        'system administrators. It is a custom-built dashboard (separate from Django\'s built-in admin) '
        'designed specifically for the Smart Classroom\'s workflow needs.')

    add_table(doc,
        ['Section', 'Capabilities'],
        [
            ['Dashboard',    'Total counts: departments, courses, teachers, students, pending approvals'],
            ['Site Admin',   'Django-admin-style model browser with live counts for all 12 apps'],
            ['Departments',  'Create, edit, list departments with course counts'],
            ['Teachers',     'Create, approve/reject/deactivate, list with status filter, and bulk approval checkboxes'],
            ['Students',     'Create, approve/reject, enroll in courses, list with status filter, and bulk approval checkboxes'],
            ['Courses',      'Create (auto-approved), edit, approve/reject teacher-created courses'],
        ],
        col_widths=[3, 13]
    )

    add_paragraph(doc,
        'The admin panel uses a clean card-based layout with count badges showing pending items '
        'that require attention. All approval actions send automatic notifications to the affected users.')

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 11: SECURITY FEATURES
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_11(doc):
    doc.add_heading('11. Security Features', level=1)

    add_paragraph(doc,
        'Security is a core concern in the Smart Classroom LMS. The system implements multiple '
        'layers of protection covering authentication, authorization, data integrity, and exam security.')

    doc.add_heading('Authentication & Authorization', level=3)
    add_bullet(doc, '**Custom User Model** (`CustomUser`) extending Django\'s `AbstractUser` with role field, account approval status, and unique student ID')
    add_bullet(doc, '**Unique Auto-Generated Student ID** — student accounts automatically receive a sequential identifier `STU-YYYYXX` on registration (e.g. STU-202601) to uniquely identify learners across academic years')
    add_bullet(doc, '**AccountStatusMiddleware** — intercepts every request; redirects unapproved users to approval page')
    add_bullet(doc, '**@login_required** decorator on all views requiring authentication')
    add_bullet(doc, '**Role-based access control** — views check `user.role` before allowing actions')
    add_bullet(doc, '**Enrollment verification** — course content views verify student enrollment before granting access')
    add_bullet(doc, '**CSRF protection** — Django\'s built-in CSRF middleware on all POST forms')

    doc.add_heading('Content Access Control', level=3)
    add_bullet(doc, '**Recorded classes** — enrollment check prevents unauthorized video access (Critical fix #12)')
    add_bullet(doc, '**Assignment submissions** — server-side deadline enforcement prevents late submissions')
    add_bullet(doc, '**Quiz/Test attempts** — unique constraints prevent multiple attempts per student')
    add_bullet(doc, '**Chat access** — students can only message teachers/peers in enrolled courses')

    doc.add_heading('Data Integrity', level=3)
    add_bullet(doc, '**Soft-delete enrollment** — `is_active=False` instead of deletion, preserving all student data')
    add_bullet(doc, '**Unique constraints** on 9 critical model relationships (see Section 6.3)')
    add_bullet(doc, '**Server-side validation** — all forms validated on the server, not just client-side')
    add_bullet(doc, '**Password validators** — Django\'s 4 built-in password validators enforced')

    doc.add_heading('Exam Security (Anti-Cheat)', level=3)
    add_paragraph(doc,
        'The anti-cheat system (detailed in Section 9) prevents students from cheating during '
        'online quizzes and coding tests. It combines fullscreen enforcement, tab/window monitoring, '
        'and automatic submission on violation threshold.')

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 12: SETUP & INSTALLATION
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_12(doc):
    doc.add_heading('12. Setup & Installation', level=1)

    add_paragraph(doc,
        'The project includes an automated setup script (`setup.py`) that handles the complete '
        'installation process including virtual environment creation, dependency installation, '
        'database migration, sample data creation, and server startup.')

    doc.add_heading('Quick Setup (Automated)', level=3)
    add_code_block(doc, """# Navigate to project directory
cd "c:\\Data Science and Gen Ai projects\\smart_classroom"

# Run automated setup (creates venv, installs deps, migrates, seeds data, starts server)
python setup.py""")

    doc.add_heading('Manual Setup (Step-by-Step)', level=3)
    add_code_block(doc, """# 1. Navigate to project
cd "c:\\Data Science and Gen Ai projects\\smart_classroom"

# 2. Create & activate virtual environment
python -m venv .venv
.venv\\Scripts\\activate        # Windows
# source .venv/bin/activate   # Linux/Mac

# 3. Install dependencies
pip install -r requirements.txt

# 4. Run database migrations
python manage.py migrate

# 5. Create sample data (teachers, students, courses, etc.)
python setup.py

# 6. Start the development server
python manage.py runserver

# 7. Open browser → http://localhost:8000""")

    doc.add_heading('Default Test Accounts', level=3)
    add_paragraph(doc, 'After running the setup, the following test accounts are available:')

    doc.add_heading('Teachers', level=4)
    add_table(doc,
        ['Username', 'Password', 'Role'],
        [
            ['teacher1', 'pass123', 'Teacher'],
            ['teacher2', 'pass123', 'Teacher'],
        ],
        col_widths=[5, 5, 5]
    )

    doc.add_heading('Students', level=4)
    add_table(doc,
        ['Username', 'Password', 'Role'],
        [
            ['student1', 'pass123', 'Student'],
            ['student2', 'pass123', 'Student'],
            ['student3', 'pass123', 'Student'],
            ['student4', 'pass123', 'Student'],
            ['student5', 'pass123', 'Student'],
        ],
        col_widths=[5, 5, 5]
    )

    doc.add_heading('Pre-Created Demo Data', level=3)
    add_table(doc,
        ['Data Type', 'Count', 'Details'],
        [
            ['Courses',          '3', 'Python Programming, Web Development, Data Science'],
            ['Assignments',      '2 per course', 'With file attachments'],
            ['Quizzes',          '1 per course', 'With MCQ questions'],
            ['Materials',        '2 per course', 'Document + link types'],
            ['Recorded Classes', '1 per course', 'Video files'],
            ['Projects',         '1 per course', 'With descriptions and deadlines'],
            ['Certificates',     'Generated',    'For enrolled students'],
            ['Report Cards',     'Generated',    'With aggregated scores'],
        ],
        col_widths=[4, 3, 9]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 13: URL REFERENCE MAP
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_13(doc):
    doc.add_heading('13. Complete URL Reference Map', level=1)

    add_paragraph(doc,
        'The Smart Classroom LMS exposes 105+ URL endpoints organized across 12 Django apps. '
        'This reference map documents every accessible URL in the system.')

    # Users
    doc.add_heading('Users (11 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/users/register/',          'Registration page (select role, fill profile)'],
            ['/users/login/',             'Login page (username + password)'],
            ['/users/logout/',            'Logout (redirects to login)'],
            ['/users/dashboard/',         'Role-based dashboard (different per role)'],
            ['/users/profile/',           'Edit profile (name, bio, picture)'],
            ['/users/change-password/',   'Change password form'],
            ['/users/leaderboard/',       'Global/course leaderboard rankings'],
            ['/users/report-card/',       'Student report card view'],
            ['/users/report-card/download/', 'Download report card as PDF'],
            ['/users/approval-status/',   'Pending approval status page'],
            ['/users/notifications/',     'Notification inbox'],
        ],
        col_widths=[6.5, 9.5]
    )

    # Admin Panel
    doc.add_heading('Admin Panel (16 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/users/admin-panel/',                       'Admin dashboard'],
            ['/users/admin-panel/site-admin/',             'Django model browser'],
            ['/users/admin-panel/departments/',             'Department list'],
            ['/users/admin-panel/departments/create/',      'Create department'],
            ['/users/admin-panel/departments/<pk>/edit/',   'Edit department'],
            ['/users/admin-panel/teachers/',                'Teacher list'],
            ['/users/admin-panel/teachers/create/',         'Create teacher'],
            ['/users/admin-panel/teachers/<pk>/approve/',   'Approve teacher'],
            ['/users/admin-panel/teachers/<pk>/reject/',    'Reject teacher'],
            ['/users/admin-panel/teachers/<pk>/deactivate/','Deactivate teacher'],
            ['/users/admin-panel/students/',                'Student list'],
            ['/users/admin-panel/students/create/',         'Create student'],
            ['/users/admin-panel/students/<pk>/approve/',   'Approve student'],
            ['/users/admin-panel/students/<pk>/reject/',    'Reject student'],
            ['/users/admin-panel/students/<pk>/enroll/',    'Enroll student in course'],
            ['/users/admin-panel/courses/',                 'Course management list'],
        ],
        col_widths=[8, 8]
    )

    # Courses
    doc.add_heading('Courses (14 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/courses/',                          'Department list (course browser)'],
            ['/courses/create/',                   'Create new course'],
            ['/courses/department/<dept_pk>/',      'Courses in department'],
            ['/courses/<pk>/',                      'Course detail page'],
            ['/courses/<pk>/about/',                'Course about/info page'],
            ['/courses/<pk>/edit/',                 'Edit course details'],
            ['/courses/<pk>/delete/',               'Delete course'],
            ['/courses/<pk>/enroll/',               'Request enrollment'],
            ['/courses/<pk>/unenroll/',             'Cancel/withdraw enrollment'],
            ['/courses/enrollment/<pk>/approve/',   'Approve enrollment request'],
            ['/courses/enrollment/<pk>/reject/',    'Reject enrollment request'],
            ['/courses/<id>/students/',             'Get enrolled students (JSON API)'],
            ['/courses/<pk>/sessions/create/',      'Create virtual session'],
            ['/courses/<pk>/ai-tutor/',             'AI Tutor chatbot interface'],
        ],
        col_widths=[7, 9]
    )

    # Assignments
    doc.add_heading('Assignments (9 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/assignments/course/<course_pk>/create/', 'Create assignment for course'],
            ['/assignments/<pk>/',                       'Assignment detail page'],
            ['/assignments/<pk>/edit/',                   'Edit assignment'],
            ['/assignments/<pk>/submit/',                 'Submit assignment (file/GitHub)'],
            ['/assignments/<pk>/delete/',                 'Delete assignment'],
            ['/assignments/submission/<pk>/grade/',       'Grade student submission'],
            ['/assignments/submission/<pk>/delete/',      'Delete submission'],
            ['/assignments/course/<pk>/export/csv/',     'Export marks as CSV'],
            ['/assignments/course/<pk>/export/excel/',   'Export marks as Excel'],
        ],
        col_widths=[7.5, 8.5]
    )

    # Attendance, Materials, Quiz, Tests, Projects, Chat, Certs, Reports, Recorded
    doc.add_heading('Attendance (5 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/attendance/course/<pk>/',          'Attendance session list'],
            ['/attendance/course/<pk>/create/',    'Create attendance session'],
            ['/attendance/session/<pk>/',          'Take/mark attendance'],
            ['/attendance/session/<pk>/mark/',     'Student self-mark attendance'],
            ['/attendance/course/<pk>/student/',   'Student attendance view'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Materials (3 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/materials/course/<pk>/upload/', 'Upload study material'],
            ['/materials/<pk>/download/',       'Download material file'],
            ['/materials/<pk>/delete/',         'Delete material'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Quizzes (10 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/quiz/course/<pk>/',          'Quiz list for course'],
            ['/quiz/course/<pk>/create/',    'Create new quiz'],
            ['/quiz/<pk>/edit/',             'Edit quiz details'],
            ['/quiz/<pk>/questions/',         'Add/manage questions'],
            ['/quiz/<pk>/attempt/',           'Take quiz (student)'],
            ['/quiz/result/<pk>/',            'View quiz result'],
            ['/quiz/<pk>/submissions/',       'All quiz submissions (teacher)'],
            ['/quiz/<pk>/students/',          'Students who attempted'],
            ['/quiz/<pk>/delete/',            'Delete quiz'],
            ['/quiz/question/<pk>/delete/',   'Delete individual question'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Online Tests (12 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/tests/',                         'Test list (all)'],
            ['/tests/course/<pk>/',              'Tests by course'],
            ['/tests/create/',                   'Create test (no course)'],
            ['/tests/course/<pk>/create/',        'Create test for specific course'],
            ['/tests/<id>/',                      'Test detail page'],
            ['/tests/<id>/edit/',                 'Edit test'],
            ['/tests/<id>/delete/',               'Delete test'],
            ['/tests/<id>/take/',                 'Take test (coding IDE)'],
            ['/tests/<id>/question/add/',         'Add coding question'],
            ['/tests/question/<id>/cases/add/',   'Add test cases to question'],
            ['/tests/<id>/responses/',             'All student responses'],
            ['/tests/<id>/responses/<id>/',        'Individual response detail'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Projects (4 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/projects/',             'Project list'],
            ['/projects/submit/',      'Submit project (GitHub/ZIP)'],
            ['/projects/<pk>/',         'Project detail + grading'],
            ['/projects/<pk>/delete/',  'Delete project'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Chat (6 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/chat/',                       'Chat inbox (conversation list)'],
            ['/chat/<user_id>/',              'Direct message chat room'],
            ['/chat/send/<user_id>/',          'Send DM (AJAX POST)'],
            ['/chat/fetch/<user_id>/',         'Fetch new messages (AJAX GET)'],
            ['/chat/group/<course_id>/',        'Course group chat room'],
            ['/chat/group/fetch/<course_id>/',  'Fetch group messages (AJAX GET)'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Certificates (6 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/certificates/course/<pk>/issue/',  'Issue certificates to students'],
            ['/certificates/my/',                  'My certificates list'],
            ['/certificates/<pk>/view/',            'View certificate details'],
            ['/certificates/<pk>/exact/',           'Detailed certificate view'],
            ['/certificates/<pk>/download/',        'Download certificate as PDF'],
            ['/certificates/<pk>/exact/download/',  'Download detailed PDF'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Reports (4 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/reports/',            'Reports home page'],
            ['/reports/manage/',     'Report management list (teacher)'],
            ['/reports/<pk>/',        'Individual report detail'],
            ['/reports/<pk>/edit/',    'Edit report card'],
        ],
        col_widths=[7, 9]
    )

    doc.add_heading('Recorded Classes (5 URLs)', level=3)
    add_table(doc,
        ['URL', 'Description'],
        [
            ['/recorded-classes/',              'Recording list'],
            ['/recorded-classes/upload/',        'Upload new recording'],
            ['/recorded-classes/<pk>/edit/',      'Edit recording details'],
            ['/recorded-classes/<pk>/delete/',    'Delete recording'],
            ['/recorded-classes/<pk>/',            'Watch recording (video player)'],
        ],
        col_widths=[7, 9]
    )

    # Final count summary
    doc.add_heading('URL Count Summary', level=3)
    add_table(doc,
        ['Module', 'URL Count'],
        [
            ['Users',           '11'],
            ['Admin Panel',     '16'],
            ['Courses',         '14'],
            ['Assignments',     '9'],
            ['Attendance',      '5'],
            ['Materials',       '3'],
            ['Quizzes',         '10'],
            ['Online Tests',    '12'],
            ['Projects',        '4'],
            ['Chat',            '6'],
            ['Certificates',    '6'],
            ['Reports',         '4'],
            ['Recorded Classes','5'],
            ['TOTAL',           '105+'],
        ],
        col_widths=[6, 5]
    )

    add_page_break(doc)


# ═══════════════════════════════════════════════════════════════════════════════
#  SECTION 14: DEVELOPMENT HISTORY & BUG FIXES
# ═══════════════════════════════════════════════════════════════════════════════

def build_section_14(doc):
    doc.add_heading('14. Development History & Bug Fixes', level=1)

    add_paragraph(doc,
        'During the development of Smart Classroom LMS, 37 issues were identified, diagnosed, '
        'and fixed across the entire application. This section documents every issue with its '
        'impact, priority, and the solution applied. All issues are now resolved.')

    doc.add_heading('Priority Distribution', level=3)
    add_table(doc,
        ['Priority', 'Count', 'Description'],
        [
            ['🔴 Critical', '5', 'Security breaches, core functionality broken'],
            ['High',         '10', 'Major feature failures, data integrity issues'],
            ['Medium',       '15', 'Feature gaps, UX problems, missing functionality'],
            ['Low',          '7',  'Visual glitches, minor polish items'],
        ],
        col_widths=[3, 2, 11]
    )

    doc.add_heading('Complete Issue Summary', level=3)

    issues = [
        ['1',  'Course cards uneven layout (single course)',    'Courses UI',       'Medium',     '✅ Fixed'],
        ['2',  'Course description bullets not rendering',      'Courses UI',       'Medium',     '✅ Fixed'],
        ['3',  'Button misalignment in course cards',           'Courses UI',       'Low',        '✅ Fixed'],
        ['4',  'Long descriptions overlapping cards',           'Courses UI',       'Medium',     '✅ Fixed'],
        ['5',  'Test time lock not working / uneven desc',      'Tests App',        'High',       '✅ Fixed'],
        ['6',  'Teacher dashboard charts not displaying',       'Dashboard',        'High',       '✅ Fixed'],
        ['7',  'Pending work false positives/negatives',        'Dashboard',        'High',       '✅ Fixed'],
        ['8',  'Projects no deadline → wrong pending status',   'Projects',         'High',       '✅ Fixed'],
        ['9',  'Chat missing emoji, file upload, poor UI',      'Chat App',         'Medium',     '✅ Fixed'],
        ['10', 'Specific students misalignment in edit',        'Assignments',      'Medium',     '✅ Fixed'],
        ['11', 'Course desc not rendering on confirm page',     'Courses',          'Low',        '✅ Fixed'],
        ['12', 'Unauthorized recorded class access',            'Recorded Classes', '🔴 Critical','✅ Fixed'],
        ['13', 'Re-enrollment no confirmation',                 'Courses',          'Medium',     '✅ Fixed'],
        ['14', 'Student data lost on un-enroll/re-enroll',      'Courses',          'High',       '✅ Fixed'],
        ['15', 'Digital badges not created on cert issue',      'Certificates',     'Medium',     '✅ Fixed'],
        ['16', 'Year label overlapping badge section',          'Profile UI',       'Low',        '✅ Fixed'],
        ['17', 'Badge download not available',                  'Profile',          'Low',        '✅ Fixed'],
        ['18', 'Edit assignment error (specific students)',     'Assignments',      '🔴 Critical','✅ Fixed'],
        ['19', 'Teacher cannot upload recorded session',        'Recorded Classes', '🔴 Critical','✅ Fixed'],
        ['20', 'Test creation "no students" error',             'Tests App',        'High',       '✅ Fixed'],
        ['21', 'Teacher nav link overflow',                     'UI / Navbar',      'Low',        '✅ Fixed'],
        ['22', 'Teacher missing pending work section',          'Dashboard',        'Medium',     '✅ Fixed'],
        ['23', 'Tests list page not scrollable',                'Tests UI',         'Medium',     '✅ Fixed'],
        ['24', 'Excel download missing weekly data',            'Reports',          'Medium',     '✅ Fixed'],
        ['25', 'Unwanted orphan files in project',              'Cleanup',          'Low',        '✅ Fixed'],
        ['26', 'No automated setup script',                     'DevOps',           'Medium',     '✅ Fixed'],
        ['27', 'No scrollbars on list pages',                   'UI/UX',            'Low',        '✅ Fixed'],
        ['28', 'No comprehensive test suite',                   'QA / Testing',     'High',       '✅ Fixed'],
        ['29', '3 test cases failing in test suite',            'QA',               'High',       '✅ Fixed'],
        ['30', 'No live demo link without deployment',          'Demo/DevOps',      'Medium',     '✅ Advised'],
        ['31', 'Students switching tabs during exams',          'Security/Exam',    '🔴 Critical','✅ Fixed'],
        ['32', 'Auto-submit not triggering after warnings',     'Security/Exam',    '🔴 Critical','✅ Fixed'],
        ['33', 'Browser X button bypasses fullscreen',          'Security/Exam',    'High',       '✅ Fixed'],
        ['34', 'Sidebar/navbar visible during exam',            'Security/Exam',    'High',       '✅ Fixed'],
        ['35', 'Half-screen window blur not detected',          'Security/Exam',    'High',       '✅ Fixed'],
        ['36', 'Exam instructions overlay wrong alignment',     'UI / Exam',        'Low',        '✅ Fixed'],
        ['37', 'Test status stuck on "Active" after deadline',  'Tests App',        'Medium',     '✅ Fixed'],
    ]

    add_table(doc,
        ['#', 'Issue', 'Module', 'Priority', 'Status'],
        issues,
        col_widths=[0.8, 7, 3, 2.2, 1.5]
    )

    # Detailed breakdown of critical issues
    doc.add_heading('Critical Issues — Detailed Solutions', level=3)

    doc.add_heading('#12: Unauthorized Recorded Class Access', level=4)
    add_paragraph(doc,
        '**Problem**: Students not enrolled in a course could access recorded class videos by '
        'navigating directly to the URL — a complete security bypass exposing content meant to be restricted.')
    add_paragraph(doc,
        '**Solution**: Added `@login_required` + enrollment check in the RecordedClass view. '
        'The view now checks if `request.user` is enrolled in the course linked to the video '
        'and redirects to the course list with an error message if not enrolled. Teacher access '
        'is always permitted.')

    doc.add_heading('#18: Edit Assignment Error (Specific Students)', level=4)
    add_paragraph(doc,
        '**Problem**: When a teacher selected "Specific Students" for an assignment and tried to '
        'save changes, a server 500 error occurred — core functionality broken.')
    add_paragraph(doc,
        '**Solution**: The M2M (many-to-many) relationship for `specific_students` was not being '
        'properly set after the form save. Fixed with `form.save_m2m()` call in the edit view.')

    doc.add_heading('#19: Teacher Cannot Upload Recorded Session', level=4)
    add_paragraph(doc,
        '**Problem**: Upload failed silently or threw an error — no recorded content management possible.')
    add_paragraph(doc,
        '**Solution**: Fixed `MEDIA_ROOT`/`MEDIA_URL` settings, updated file upload view for large files, '
        'added `enctype="multipart/form-data"` to the upload form, and checked file size limits.')

    doc.add_heading('#31-32: Anti-Cheat System Failures', level=4)
    add_paragraph(doc,
        '**Problem**: Students could freely switch tabs/windows during exams; the warning counter '
        'didn\'t trigger auto-submit after 2 warnings.')
    add_paragraph(doc,
        '**Solution**: Implemented comprehensive anti-cheat: fullscreen lock with `visibilitychange` '
        'and `blur` event listeners, proper violation counter with `setTimeout` fallback for cooldown, '
        'and forced `form.submit()` on 3rd violation with answer auto-save via hidden fields.')

    # Final State
    doc.add_heading('Final System State', level=3)
    add_paragraph(doc, 'After all fixes, every module of the Smart Classroom LMS is fully functional:')

    add_table(doc,
        ['Module', 'Status'],
        [
            ['User Authentication (Teacher/Student)',  '✅ Working'],
            ['Course Management',                      '✅ Working'],
            ['Assignment Management',                  '✅ Working'],
            ['Quiz System',                            '✅ Working'],
            ['Tests System (with anti-cheat)',          '✅ Working'],
            ['Projects Tracking',                      '✅ Working'],
            ['Recorded Classes (access-controlled)',    '✅ Working'],
            ['Chat / Messaging',                       '✅ Working'],
            ['Attendance Tracking',                    '✅ Working'],
            ['Certificates & Digital Badges',          '✅ Working'],
            ['Reports & Analytics (Excel download)',   '✅ Working'],
            ['Teacher Dashboard with Charts',          '✅ Working'],
            ['Student Dashboard with Pending Work',    '✅ Working'],
            ['Anti-Cheat Exam System',                 '✅ Working'],
            ['Comprehensive Test Suite',               '✅ Working'],
            ['Automated Setup Script',                 '✅ Working'],
        ],
        col_widths=[8, 3]
    )

    # Final Summary
    add_divider(doc)
    doc.add_heading('Final Summary', level=3)
    add_table(doc,
        ['Metric', 'Count'],
        [
            ['Django Apps',             '12'],
            ['Database Models',         '24'],
            ['URL Endpoints',           '105+'],
            ['Templates',              '50+'],
            ['User Roles',             '3 (Admin, Teacher, Student)'],
            ['Approval Workflows',     '3 (User, Course, Enrollment)'],
            ['AI Integration Points',  '2 (AI Tutor, Coding Evaluation)'],
            ['Export Formats',         '3 (CSV, Excel, PDF)'],
            ['Chat Types',            '2 (DM + Group)'],
            ['Assessment Types',      '4 (Assignment, Quiz, Coding Test, Project)'],
            ['Certificate Types',     '4 (Completion, Excellence, Topper, Participation)'],
            ['Bugs Fixed',            '37 (5 Critical, 10 High, 15 Medium, 7 Low)'],
        ],
        col_widths=[5, 8]
    )

    add_divider(doc)

    # Document footer info
    add_paragraph(doc, '', space_after=20)
    add_paragraph(doc, 'Document prepared by: Yaswanth Munagoti',
                  italic=True, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, 'Project: Smart Classroom LMS',
                  italic=True, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_paragraph(doc, 'Last Updated: June 2026',
                  italic=True, color=COLOR_GRAY,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)


# ═══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════════

def main():
    print("[DOC] Generating Smart Classroom Complete Documentation...")
    print(f"   Output: {OUTPUT_FILE}")
    print(f"   Images: {IMAGES_DIR}")
    print()

    # Verify images exist
    expected_images = [f'diagram_{i}.png' for i in range(1, 22)] + ['complete_database_schema.png']
    missing = [img for img in expected_images if not os.path.exists(os.path.join(IMAGES_DIR, img))]
    if missing:
        print(f"[WARN] Missing images: {missing}")
        print("   Continuing -- missing images will show placeholder text.")
    else:
        print(f"[OK] All {len(expected_images)} diagram images found.")

    # Create document
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)       # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Setup styles
    setup_styles(doc)

    # Build all sections
    print("   Building cover page...")
    build_cover_page(doc)

    print("   Building table of contents...")
    build_toc(doc)

    print("   Building Section 1: Project Overview...")
    build_section_1(doc)

    print("   Building Section 2: Technology Stack...")
    build_section_2(doc)

    print("   Building Section 3: System Architecture...")
    build_section_3(doc)

    print("   Building Section 4: User Roles & Permissions...")
    build_section_4(doc)

    print("   Building Section 5: Approval Workflows...")
    build_section_5(doc)

    print("   Building Section 6: Database Schema (7 modules)...")
    build_section_6(doc)

    print("   Building Section 7: Application Flows (12 flows)...")
    build_section_7(doc)

    print("   Building Section 8: Step-by-Step Walkthrough...")
    build_section_8(doc)

    print("   Building Section 9: AI Integration...")
    build_section_9(doc)

    print("   Building Section 10: Admin Panel...")
    build_section_10(doc)

    print("   Building Section 11: Security Features...")
    build_section_11(doc)

    print("   Building Section 12: Setup & Installation...")
    build_section_12(doc)

    print("   Building Section 13: URL Reference Map...")
    build_section_13(doc)

    print("   Building Section 14: Development History (37 issues)...")
    build_section_14(doc)

    # Save
    print()
    try:
        doc.save(OUTPUT_FILE)
        file_size = os.path.getsize(OUTPUT_FILE)
        print(f"[OK] Documentation saved to: {OUTPUT_FILE}")
        print(f"   File size: {file_size / 1024:.1f} KB ({file_size / (1024*1024):.1f} MB)")
    except PermissionError:
        alt = OUTPUT_FILE.replace('.docx', '_v3.docx')
        print(f"[WARN] Permission denied (file may be open in Word). Saving as: {alt}")
        doc.save(alt)
        file_size = os.path.getsize(alt)
        print(f"[OK] Documentation saved to: {alt}")
        print(f"   File size: {file_size / 1024:.1f} KB ({file_size / (1024*1024):.1f} MB)")

    print("\n[DONE] Open the file in Microsoft Word and update the Table of Contents.")


if __name__ == '__main__':
    main()
